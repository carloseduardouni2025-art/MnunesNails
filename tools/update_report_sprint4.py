from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


DOCX_PATH = "MNunesNails_relatorio_sprint4_corrigido.docx"


def set_text(paragraph, text):
    paragraph.clear()
    if text:
        paragraph.add_run(text)


def insert_after(paragraph, text, style_name=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style_name:
        new_paragraph.style = style_name
    new_paragraph.add_run(text)
    return new_paragraph


def find_paragraph(doc, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def find_paragraph_after(doc, text, after_text=None):
    after_seen = after_text is None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == after_text:
            after_seen = True
        if after_seen and paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found after marker: {text} / {after_text}")


def replace_block_after_heading(doc, heading, texts, stop_headings=None, style_name="Texto", after_text=None):
    heading_paragraph = find_paragraph_after(doc, heading, after_text=after_text)
    stop_headings = set(stop_headings or [])
    current = heading_paragraph._p.getnext()
    first = True
    last_paragraph = heading_paragraph

    while current is not None:
        paragraph = None
        for candidate in doc.paragraphs:
            if candidate._p is current:
                paragraph = candidate
                break
        if paragraph is None:
            break
        if paragraph.text.strip() in stop_headings:
            break
        if paragraph.text.strip() or first:
            if texts:
                set_text(paragraph, texts.pop(0))
                paragraph.style = style_name
                last_paragraph = paragraph
                first = False
            else:
                set_text(paragraph, "")
        current = current.getnext()

    for text in texts:
        last_paragraph = insert_after(last_paragraph, text, style_name)


doc = Document(DOCX_PATH)

# 6. Proposta de solucao - alinha com o produto realmente desenvolvido.
replace_block_after_heading(
    doc,
    "Proposta de Solução",
    [
        "A proposta de solução consiste no desenvolvimento do MNunesNails, uma aplicação web para digitalizar o processo de agendamento do Studio MNunesnails. O sistema substitui o controle manual por uma agenda online, permitindo que clientes visualizem serviços, escolham datas e horários disponíveis e confirmem o atendimento de forma organizada.",
        "Visão geral da Solução",
        "A solução foi implementada como um sistema web responsivo, acessível pelo navegador, com páginas públicas para clientes e uma área administrativa protegida. O front-end utiliza HTML, CSS e JavaScript, enquanto o back-end foi desenvolvido em Python, com rotas de API para autenticação, usuários, serviços, agendamentos e disponibilidade. Os dados são armazenados em SQLite.",
        "Funcionalidades Principais",
        "Cadastro e login de clientes, com armazenamento de nome, telefone, WhatsApp e credenciais de acesso.",
        "Login administrativo separado, permitindo diferenciar cliente e administrador.",
        "Agendamento online com escolha de serviço, data, horário, observações e resumo do atendimento.",
        "Agenda inteligente com horários livres, bloqueios manuais, bloqueio por dia inteiro e prevenção de conflitos de acordo com a duração dos serviços.",
        "Painel administrativo com abas para agendamentos, horários, serviços e usuários.",
        "Catálogo de serviços dinâmico, permitindo adicionar, editar, ativar, desativar e excluir serviços.",
        "Integração opcional com WhatsApp, perguntando à cliente se deseja enviar a mensagem após salvar o agendamento.",
        "Diferenciais da Solução",
        "O sistema foi ajustado para refletir automaticamente, na página inicial, as alterações feitas pelo administrador nos serviços. Também foi implementado o destaque dinâmico do serviço mais agendado, tornando a vitrine mais fiel ao uso real da aplicação.",
    ],
    stop_headings={"Gestão do Backlog do Projeto (Kanban – Trello)"},
)

# 8.1 e 8.2 - explicacao das evidencias do Trello.
print_quadro = find_paragraph(doc, "Print Quadro Trello")
insert_after(
    print_quadro,
    "O print atualizado do quadro Trello deve demonstrar a Sprint 04 em andamento ou concluída, com as atividades relacionadas à finalização do MVP. As listas devem evidenciar a evolução das tarefas de autenticação, administração de agendamentos, controle de horários, gerenciamento de serviços, ajustes visuais, integração com o index e publicação no Render.",
    "Texto",
)

print_card = find_paragraph(doc, "Print de um Card")
insert_after(
    print_card,
    "O card apresentado deve estar completo, contendo descrição da tarefa, user story, critérios de aceitação e checklist de implementação. Exemplo de card utilizado na Sprint 04: como administrador, desejo adicionar, editar, ativar, desativar e excluir serviços para que o catálogo exibido às clientes seja atualizado automaticamente. Critérios de aceitação: serviço criado aparece no index; serviço desativado deixa de aparecer para a cliente; exclusão solicita confirmação; preço permanece em real; duração permanece em minutos.",
    "Texto",
)

# 9. Arquitetura - remove planejamento antigo com React e descreve arquitetura entregue.
replace_block_after_heading(
    doc,
    "Arquitetura da Solução",
    [
        "Fluxo de Desenvolvimento:",
        "- Estruturar a interface pública com HTML, CSS e JavaScript.",
        "- Criar a página inicial com apresentação do salão, catálogo de serviços e formulário de agendamento.",
        "- Implementar calendário visual para escolha de datas futuras e seleção de horários disponíveis.",
        "- Desenvolver o servidor Python responsável por servir o site e expor as rotas de API.",
        "- Criar o banco SQLite com tabelas de usuários, administradores, sessões, serviços, agendamentos e disponibilidade de horários.",
        "- Implementar autenticação separada para cliente e administrador.",
        "- Criar rotas para cadastro, login, listagem e gerenciamento de agendamentos.",
        "- Criar rotas para controle de disponibilidade, incluindo bloqueio de horário e bloqueio de dia inteiro.",
        "- Criar rotas para gerenciamento dinâmico dos serviços oferecidos.",
        "- Integrar a página administrativa ao index, para refletir automaticamente serviços ativados, desativados, adicionados ou excluídos.",
        "- Publicar o MVP no Render para validação externa.",
    ],
    stop_headings={"Repositório do Projeto (Git)"},
)

# 10. Repositorio e commits.
for paragraph in doc.paragraphs:
    if paragraph.text.strip() == "Link do repositório: link":
        set_text(paragraph, "Link do repositório: https://github.com/carloseduardouni2025-art/MnunesNails")

replace_block_after_heading(
    doc,
    "Histórico de Commits",
    [
        "a2db9b9 - Initial commit: criação inicial do repositório.",
        "4f731a6 - feat: base do projeto: estrutura inicial da aplicação web.",
        "c2c15ae - feat: pagina de agendamentos: criação da área de visualização de agendamentos.",
        "0b589b0 - feat: adm x user: separação inicial entre acesso administrativo e acesso de usuário comum.",
        "8448729 - Feat: Banco de dados e login: criação da base de dados, cadastro, login e autenticação.",
        "6f6ac71 - feat: gestão de acessos: aprimoramento das permissões e diferenciação entre cliente e administrador.",
        "1cd5000 - feat: finalização pt1: integração de funcionalidades principais do sistema.",
        "1f6f045 - Feat: adicionando configuração para render: preparação do sistema para publicação online.",
        "32e5958 - feat: redirecionamento de mensagem: ajuste do fluxo de confirmação pelo WhatsApp.",
        "af44f3d - feat: ui e ux: melhorias finais de interface, usabilidade e organização visual.",
    ],
    stop_headings={"Planejamento das Sprints"},
)

# Sprint 03 - detalhes tecnicos pedidos na avaliacao.
replace_block_after_heading(
    doc,
    "Funcionalidades implementadas",
    [
        "Durante esta sprint, foram desenvolvidas as funcionalidades de banco de dados, tela de login e gerenciamento de usuários, com foco em autenticação, organização das informações e diferenciação entre cliente e administrador.",
        "No banco de dados SQLite foram criadas as tabelas users, appointments, admins, sessions, availability_slots e services. A tabela users armazena nome, telefone, WhatsApp, telefone normalizado, senha protegida por salt e hash, além das datas de criação e atualização. A tabela appointments registra o serviço, data, horário, observações, status e vínculo com o usuário. A tabela admins guarda os dados do administrador e suas credenciais protegidas. A tabela sessions controla sessões de login. A tabela availability_slots registra horários livres ou bloqueados. A tabela services armazena nome, descrição, preço, duração e status de ativação dos serviços.",
        "A tela de login valida as informações por meio das rotas de autenticação do back-end. Para clientes, o sistema compara o telefone normalizado e a senha informada com os dados armazenados. Para administradores, a validação ocorre em rota separada, garantindo que o acesso ao painel administrativo não seja confundido com o acesso comum da cliente.",
        "O sistema diferencia cliente e administrador por meio do tipo de sessão. O cliente consegue cadastrar-se, fazer login, criar agendamentos e visualizar seus próprios horários. O administrador consegue acessar o painel de gestão, visualizar todos os agendamentos, editar dados, cancelar atendimentos, gerenciar usuários, controlar horários e administrar os serviços oferecidos.",
        "Essa estrutura aumentou a rastreabilidade dos dados e permitiu que cada perfil acessasse apenas as funcionalidades adequadas ao seu papel no sistema.",
    ],
    stop_headings={"Commits relevantes"},
    after_text="Sprint 03",
)

replace_block_after_heading(
    doc,
    "Commits relevantes",
    [
        "8448729 - Feat: Banco de dados e login. Este commit representa a criação da estrutura de persistência e das telas de autenticação, permitindo armazenar usuários, administradores, sessões e agendamentos.",
        "0b589b0 - feat: adm x user. Este commit contribuiu para a separação entre usuário comum e administrador, possibilitando fluxos de acesso diferentes dentro do sistema.",
        "6f6ac71 - feat: gestão de acessos. Este commit reforçou a diferenciação de permissões, permitindo que o administrador gerencie o sistema e que o cliente visualize apenas seus próprios agendamentos.",
    ],
    stop_headings={"Evidências"},
    after_text="Sprint 03",
)

replace_block_after_heading(
    doc,
    "Evidências",
    [
        "As evidências desta sprint devem estar diretamente relacionadas às funcionalidades de banco de dados, login e gerenciamento de usuários. A tela de login comprova a separação entre acesso de cliente e acesso administrativo. A estrutura do banco comprova a criação das tabelas necessárias para armazenar usuários, administradores, sessões, serviços, disponibilidade e agendamentos.",
        "O fluxo de cadastro armazena nome, WhatsApp, telefone normalizado e credenciais protegidas. O login valida esses dados e cria uma sessão para identificar o perfil autenticado. Quando a sessão pertence a um cliente, o sistema restringe o acesso aos próprios agendamentos. Quando a sessão pertence ao administrador, o sistema libera o painel de controle com gerenciamento completo.",
        "Essas evidências comprovam que a Sprint 03 entregou a base de autenticação, persistência e controle de acesso necessária para a evolução do MVP.",
    ],
    stop_headings={"Problemas encontrados"},
    after_text="Sprint 03",
)

# Sprint 04 - deixa entrega mais robusta e rastreavel.
replace_block_after_heading(
    doc,
    "Funcionalidades implementadas",
    [
        "Na Sprint 04, o foco foi concluir e comprovar o MVP do MNunesNails. Foram removidos elementos visuais desnecessários da página inicial, simplificando a experiência da cliente e deixando em destaque o fluxo principal de agendamento.",
        "A página administrativa foi reorganizada em abas de Agendamentos, Horários, Serviços e Usuários. Essa divisão deixou o gerenciamento mais objetivo e permitiu controlar cada área do sistema separadamente.",
        "Foi implementado um calendário visual para seleção de datas, tanto na experiência da cliente quanto na administração dos horários. O calendário permite localizar datas futuras com mais facilidade e exibe horários disponíveis de 8h às 18h.",
        "Foram criadas regras automáticas para disponibilidade, considerando a duração de cada serviço. Assim, quando um atendimento ocupa mais de um slot, os horários seguintes são bloqueados para evitar conflitos.",
        "O catálogo de serviços passou a ser dinâmico. O administrador consegue adicionar, editar, ativar, desativar e excluir serviços. As alterações são refletidas automaticamente na página inicial, centralizando os cards ativos e atualizando o seletor de serviços do agendamento.",
        "O destaque 'Mais pedido' deixou de ser fixo e passou a considerar a quantidade real de agendamentos não cancelados de cada serviço.",
        "O fluxo de WhatsApp foi ajustado para perguntar à cliente se deseja enviar a mensagem após salvar o agendamento, evitando redirecionamento automático.",
    ],
    stop_headings={"Commits relevantes"},
    after_text="Sprint 04",
)

replace_block_after_heading(
    doc,
    "Commits relevantes",
    [
        "1cd5000 - feat: finalização pt1. Integração das funcionalidades principais necessárias para consolidar o MVP.",
        "1f6f045 - Feat: adicionando configuração para render. Preparação do projeto para publicação na plataforma Render.",
        "32e5958 - feat: redirecionamento de mensagem. Ajuste do comportamento do WhatsApp, substituindo o redirecionamento automático por confirmação da cliente.",
        "af44f3d - feat: ui e ux. Melhorias de interface e usabilidade, incluindo ajustes visuais, organização das telas e refinamento da experiência administrativa e da cliente.",
    ],
    stop_headings={"Evidências"},
    after_text="Sprint 04",
)

replace_block_after_heading(
    doc,
    "Evidências",
    [
        "As evidências da Sprint 04 demonstram a finalização do MVP em funcionamento. A página inicial comprova o catálogo de serviços, o formulário de agendamento, o calendário visual, os horários disponíveis e o resumo do atendimento.",
        "A área administrativa comprova a divisão em abas, a listagem e edição de agendamentos, o controle de horários, o bloqueio de dia inteiro, o gerenciamento de usuários e a administração dos serviços oferecidos.",
        "A publicação em https://mnunesnails-1.onrender.com/ comprova que o sistema foi disponibilizado em ambiente externo para validação. As rotas públicas /api/services e /api/availability demonstram que os serviços e os horários são carregados dinamicamente pelo back-end.",
    ],
    stop_headings={"Problemas encontrados"},
    after_text="Sprint 04",
)

# 13. MVP Desenvolvido - refinado com base no codigo e no link publicado.
replace_block_after_heading(
    doc,
    "MVP Desenvolvido",
    [
        "O Produto Mínimo Viável desenvolvido foi o MNunesNails, disponível em https://mnunesnails-1.onrender.com/, uma aplicação web para agendamento online de serviços de manicure e gerenciamento administrativo do Studio MNunesnails. O sistema permite que a cliente visualize os serviços oferecidos, escolha uma data pelo calendário, selecione um horário disponível, preencha seus dados e confirme o agendamento com opção de envio das informações pelo WhatsApp.",
        "As funcionalidades entregues incluem página inicial com apresentação do salão, catálogo dinâmico de serviços, formulário de agendamento, calendário visual para escolha de datas futuras, exibição automática dos horários disponíveis, resumo do agendamento, confirmação opcional via WhatsApp, tela de login, cadastro de cliente e página para acompanhamento dos próprios agendamentos.",
        "Na área administrativa, o responsável pelo salão consegue visualizar e editar agendamentos, filtrar registros por data e status, cancelar atendimentos, controlar horários livres ou bloqueados, bloquear um dia inteiro, gerenciar usuários cadastrados e administrar os serviços oferecidos. Os serviços podem ser adicionados, editados, ativados, desativados ou excluídos, e essas mudanças são refletidas automaticamente na página principal.",
        "O sistema calcula a disponibilidade considerando a duração dos serviços cadastrados. Dessa forma, quando um agendamento ocupa mais tempo, os horários seguintes são bloqueados automaticamente, evitando sobreposição entre atendimentos. Horários confirmados somente são liberados novamente quando o agendamento é cancelado.",
        "As tecnologias utilizadas no MVP foram HTML, CSS e JavaScript no front-end, com interface responsiva e temática voltada para salão de manicure. No back-end, foi utilizado Python com rotas de API próprias para autenticação, serviços, usuários, agendamentos e disponibilidade de horários. O armazenamento dos dados foi feito com SQLite, por meio do arquivo mnunesnails.db. O projeto também recebeu dependências de Express, TypeORM, PostgreSQL e dotenv para preparação técnica, mas a versão publicada analisada utiliza principalmente Python e SQLite.",
        "Como limitações do MVP, o sistema ainda pode evoluir em segurança, autenticação, controle avançado de permissões, proteção de dados, validação contra acessos simultâneos e uso de banco de dados mais robusto em produção. Também é recomendado substituir credenciais padrão, ampliar testes de responsividade e melhorar a infraestrutura definitiva de hospedagem.",
    ],
    stop_headings={"Entrega do MVP"},
)

replace_block_after_heading(
    doc,
    "Implantação do Sistema",
    [
        "O MVP foi implantado na plataforma Render, permitindo acesso externo ao sistema pelo endereço https://mnunesnails-1.onrender.com/. A implantação exigiu a preparação do projeto para execução em ambiente online, mantendo o servidor responsável por entregar as páginas HTML, arquivos CSS e JavaScript, além das rotas de API consumidas pelo front-end.",
        "A publicação possibilitou validar o fluxo principal fora do ambiente local: carregamento da página inicial, consulta dinâmica dos serviços, consulta de disponibilidade de horários e acesso às telas do sistema.",
    ],
    stop_headings={"Evidências da Entrega Presencial"},
)

replace_block_after_heading(
    doc,
    "Evidências da Entrega Presencial",
    [
        "Como evidências técnicas da entrega do MVP, foram registrados o link público do sistema publicado no Render, a página inicial com o fluxo de agendamento, a área administrativa com abas de gerenciamento, as rotas de serviços e disponibilidade funcionando e o histórico de commits do GitHub.",
        "As fotos da apresentação presencial e da entrega formal à instituição devem ser anexadas pelo aluno após a realização do encontro com o responsável institucional, pois esse material depende de registro externo ao código desenvolvido.",
    ],
    stop_headings={"Relato da Instituição"},
)

replace_block_after_heading(
    doc,
    "Relato da Instituição",
    [
        "O relato formal da instituição parceira deve ser coletado com a responsável pelo Studio MNunesnails após a apresentação do MVP. A validação esperada deve considerar se o sistema facilita o agendamento, reduz a dependência de controle manual e melhora a organização dos horários do salão.",
    ],
    stop_headings={"Impacto Social"},
)

replace_block_after_heading(
    doc,
    "Impacto Social",
    [
        "O MNunesNails contribui para a instituição ao organizar digitalmente o processo de agendamento, reduzindo a dependência de anotações manuais e diminuindo o risco de conflitos de horários. Com a agenda online, a cliente consegue consultar serviços, escolher uma data e selecionar um horário disponível com mais autonomia.",
        "Para o Studio MNunesnails, o sistema melhora a visibilidade dos serviços oferecidos, padroniza os dados dos atendimentos e facilita a gestão administrativa. A responsável pelo salão passa a ter uma visão mais clara dos agendamentos, dos horários bloqueados, dos usuários cadastrados e dos serviços ativos.",
        "Para as clientes, o benefício principal está na praticidade. O processo de agendamento fica mais simples, rápido e acessível pelo navegador, com possibilidade de confirmação pelo WhatsApp. Isso fortalece a comunicação entre cliente e estabelecimento e melhora a experiência de atendimento.",
    ],
    stop_headings={"Reflexão Crítica"},
)

replace_block_after_heading(
    doc,
    "Reflexão Crítica",
    [
        "O desenvolvimento do MNunesNails proporcionou aprendizado técnico em front-end, back-end, banco de dados, autenticação, integração entre telas e publicação de aplicação web. Ao longo das sprints, foi possível perceber a importância de documentar cada etapa com evidências claras e commits específicos.",
        "Entre os principais desafios estiveram a organização das regras de horários, a diferenciação entre cliente e administrador, a atualização dinâmica dos serviços e a prevenção de conflitos entre agendamentos. Esses pontos exigiram ajustes sucessivos e maior atenção à lógica de negócio.",
        "A experiência também demonstrou que um MVP deve priorizar o fluxo essencial do usuário. No caso do MNunesNails, a entrega mais importante foi garantir que a cliente consiga agendar online e que o administrador consiga controlar serviços, horários, usuários e agendamentos de forma prática.",
    ],
    stop_headings={"Próximos Passos"},
)

replace_block_after_heading(
    doc,
    "Próximos Passos",
    [
        "Como próximos passos, recomenda-se fortalecer a segurança da autenticação, substituir credenciais padrão, preparar o banco de dados para ambiente definitivo de produção e melhorar o tratamento de acessos simultâneos.",
        "Também é importante ampliar os testes em dispositivos móveis, revisar responsividade, criar relatórios simples para serviços mais agendados e horários de maior procura, além de melhorar mensagens de confirmação para ações administrativas.",
        "Outra evolução possível é integrar notificações automáticas, histórico detalhado de alterações e recursos de confirmação ou lembrete para clientes, tornando o sistema mais completo para uso contínuo no salão.",
    ],
    stop_headings={"Apêndices"},
)

replace_block_after_heading(
    doc,
    "Apêndices",
    [
        "Quadro de Autoavaliação do Projeto",
        "O projeto possui evidências no repositório GitHub, no histórico de commits, no site publicado no Render e nas capturas de tela inseridas ao longo do relatório.",
        "Itens concluídos: levantamento do problema, proposta de solução, organização por sprints, desenvolvimento do MVP, autenticação, banco de dados, painel administrativo, gerenciamento de serviços, calendário, controle de horários, integração com WhatsApp e publicação online.",
        "Video de Apresentação do Pitch",
        "Link do Drive de acesso ao vídeo: inserir o link após a gravação e envio do pitch final.",
        "Produções adicionais do aluno",
        "Foram produzidos o código-fonte do sistema MNunesNails, o banco de dados SQLite, a interface pública de agendamento, a área administrativa, as rotas de API, a publicação no Render e a documentação incremental do desenvolvimento.",
    ],
    stop_headings=set(),
)

for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if text in {
        "Orientação ao estudante",
        "Abaixo da tabela, caso o sistema tenha sido implantado, relatar como aconteceu a implantação. Caso não tenha sido implantado, justificar o motivo.",
        "*Essas orientações devem ser removidas no envio do documento.",
    }:
        set_text(paragraph, "")

appendix_tail_seen = False
for paragraph in doc.paragraphs:
    if paragraph.text.strip().startswith("Foram produzidos o código-fonte do sistema MNunesNails"):
        appendix_tail_seen = True
        continue
    if appendix_tail_seen:
        set_text(paragraph, "")

doc.save(DOCX_PATH)
